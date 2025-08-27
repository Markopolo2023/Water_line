import sqlite3
import os
import re
from datetime import datetime
from docx import Document
import pdfplumber
import logging
import sys
import logging as lib_logging  # Improvement: For suppressing library logs

# Improvement: Suppress pdfplumber debug logs
lib_logging.getLogger('pdfplumber').setLevel(lib_logging.WARNING)

def setup_logging(log_dir):
    """Configure logging to file and console with UTF-8 encoding."""
    log_file = os.path.join(log_dir, 'report_processor.log')
    logging.basicConfig(
        level=logging.DEBUG,
        format='%(asctime)s - %(levelname)s - %(message)s',
        handlers=[
            logging.FileHandler(log_file, encoding='utf-8'),
            logging.StreamHandler(sys.stdout)
        ]
    )
    logger = logging.getLogger(__name__)
    return logger

def process_reports(reports_path, db_path):
    """Process reports from the specified directory and store data in the SQLite database."""
    # Set console encoding to UTF-8 to handle special characters
    if sys.platform == "win32":
        import io
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

    logger = setup_logging(os.path.dirname(db_path))
    logger.info(f"Starting processing of files in directory: {reports_path}")
    logger.debug(f"Database path: {db_path}")

    # Ensure the reports directory exists
    if not os.path.exists(reports_path):
        logger.info(f"Creating reports directory: {reports_path}")
        os.makedirs(reports_path)

    # Connect to SQLite database
    try:
        conn = sqlite3.connect(db_path)
        cur = conn.cursor()
        logger.info(f"Successfully connected to database: {db_path}")
    except Exception as e:
        logger.error(f"Failed to connect to database {db_path}: {e}")
        raise

    # Create reports table if not exists
    try:
        cur.execute('''
        CREATE TABLE IF NOT EXISTS reports (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            filename TEXT UNIQUE,  -- Improvement: Add UNIQUE to prevent duplicate filenames
            facility TEXT,
            date DATE,
            chemist TEXT
        )
        ''')
        logger.info("Reports table created or already exists")
    except Exception as e:
        logger.error(f"Error creating reports table: {e}")
        raise

    # Create systems table if not exists
    try:
        cur.execute('''
        CREATE TABLE IF NOT EXISTS systems (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            report_id INTEGER,
            system_name TEXT,
            cond REAL,
            ph REAL,
            temp REAL,
            p_alk REAL,
            m_alk REAL,
            chloride REAL,
            hardness REAL,
            calcium REAL,
            po4 REAL,
            so3 REAL,
            mo REAL,
            no2 REAL,
            live_atp REAL,
            glycol TEXT,
            free_chlorine REAL,
            total_chlorine REAL,
            max_temp REAL,
            comment TEXT,
            UNIQUE(report_id, system_name)  -- Improvement: Prevent duplicate systems per report
        )
        ''')
        logger.info("Systems table created or already exists")
    except Exception as e:
        logger.error(f"Error creating systems table: {e}")
        raise

    conn.commit()

    # Helper function to parse date strings in various formats
    def parse_date(date_str):
        formats = ['%m/%d/%Y', '%m-%d-%Y', '%m/%d/%y', '%m-%d-%y']
        for fmt in formats:
            try:
                date = datetime.strptime(date_str, fmt).date()
                logger.debug(f"Parsed date '{date_str}' as {date} using format {fmt}")
                return date
            except ValueError:
                continue
        logger.warning(f"Could not parse date: {date_str}")
        return None

    # Process each file in the folder
    files = [f for f in os.listdir(reports_path) if f.lower().endswith(('.docx', '.pdf'))]
    logger.info(f"Found {len(files)} files to process: {files}")

    for filename in files:
        full_path = os.path.join(reports_path, filename)
        # Improvement: Skip empty files early
        if os.path.getsize(full_path) == 0:
            logger.warning(f"Skipping empty file: {filename}")
            continue
        logger.info(f"Processing file: {filename}")
        logger.debug(f"File size: {os.path.getsize(full_path)} bytes")  # Improvement: Log file size
        full_text = ''
        tables = []

        # Extract text and tables based on file type
        if filename.lower().endswith('.docx'):
            try:
                doc = Document(full_path)
                # Improvement: Force UTF-8 handling
                full_text = '\n'.join(paragraph.text.encode('utf-8', 'ignore').decode('utf-8') for paragraph in doc.paragraphs)
                logger.debug(f"Extracted {len(full_text)} characters of text from {filename}")
                for table in doc.tables:
                    table_data = [[cell.text.strip().encode('utf-8', 'ignore').decode('utf-8') if cell.text is not None else '' for cell in row.cells] for row in table.rows]
                    tables.append(table_data)
                    logger.debug(f"Extracted table with {len(table_data)} rows from {filename}")
            except Exception as e:
                logger.error(f"Error processing DOCX {filename}: {e}")
                continue
        elif filename.lower().endswith('.pdf'):
            try:
                with pdfplumber.open(full_path) as pdf:
                    # Improvement: Force UTF-8 handling
                    full_text = '\n'.join(page.extract_text() or '' for page in pdf.pages).encode('utf-8', 'ignore').decode('utf-8')
                    logger.debug(f"Extracted {len(full_text)} characters of text from {filename}")
                    for page in pdf.pages:
                        extracted_table = page.extract_table()
                        if extracted_table:
                            cleaned_table = [[cell.encode('utf-8', 'ignore').decode('utf-8') if cell is not None else '' for cell in row] for row in extracted_table]
                            tables.append(cleaned_table)
                            logger.debug(f"Extracted table with {len(cleaned_table)} rows from {filename}")
            except Exception as e:
                logger.error(f"Error processing PDF {filename}: {e}")
                continue

        if not full_text.strip():
            logger.warning(f"No text extracted from {filename}, skipping")
            continue

        logger.debug(f"Total tables extracted from {filename}: {len(tables)}")

        # Extract facility
        facility = ''
        match = re.search(r'Facility:\s*(.+?)(?:\n|$)', full_text, re.I)
        if match:
            facility = match.group(1).strip()
            logger.debug(f"Extracted facility: {facility}")
        else:
            match = re.search(r'SITE VISITATION REPORT\s*\n\s*(.+?)(?:\n|$)', full_text, re.I)
            if match:
                facility = match.group(1).strip()
                logger.debug(f"Extracted facility from SITE VISITATION REPORT: {facility}")
            else:
                logger.warning(f"No facility found in {filename}")

        # Extract date
        date_str = ''
        match = re.search(r'Date:\s*([\d/-]+)', full_text, re.I)
        if match:
            date_str = match.group(1).strip()
            logger.debug(f"Extracted date string: {date_str}")
        date = parse_date(date_str)
        date_value = date.isoformat() if date else None
        logger.debug(f"Parsed date value: {date_value}")

        # Extract chemist
        chemist = ''
        match = re.search(r'IWT Field Representative:\s*(.+?),\s*.+@illinois\.edu', full_text, re.I)
        if match:
            chemist = match.group(1).strip()
            logger.debug(f"Extracted chemist: {chemist}")
        else:
            match = re.search(r'Signature\s*\n\s*(.+?),\s*BS Ch\.E\., Senior Chemist', full_text, re.I)
            if match:
                chemist = match.group(1).strip()
                logger.debug(f"Extracted chemist from Signature: {chemist}")
            else:
                match = re.search(r'x([A-Z][a-z]+ [A-Z][a-z]+)\n\1, Field Chemist', full_text, re.I)
                if match:
                    chemist = match.group(1).strip()
                    logger.debug(f"Extracted chemist from Field Chemist: {chemist}")
                else:
                    logger.warning(f"No chemist found in {filename}")

        # Insert into reports table (Improvement: Check for existing)
        cur.execute("SELECT id FROM reports WHERE filename = ?", (filename,))
        existing = cur.fetchone()
        if existing:
            report_id = existing[0]
            logger.info(f"Skipping duplicate report for filename={filename}, existing report_id={report_id}")
            continue  # Or update if needed
        try:
            cur.execute('''
            INSERT INTO reports (filename, facility, date, chemist)
            VALUES (?, ?, ?, ?)
            ''', (filename, facility, date_value, chemist))
            report_id = cur.lastrowid
            conn.commit()
            logger.info(f"Inserted into reports table: filename={filename}, facility={facility}, date={date_value}, chemist={chemist}, report_id={report_id}")
        except Exception as e:
            logger.error(f"Error inserting into reports table for {filename}: {e}")
            continue

        # Extract systems data from tables
        for table_idx, table in enumerate(tables):
            if len(table) < 2:
                logger.debug(f"Table {table_idx} in {filename} has fewer than 2 rows, skipping")
                continue
            header_row = table[0]
            header_row = [str(h) if h is not None else '' for h in header_row]
            header_str = ''.join(header_row).lower().replace('âµ', 'µ').replace('â°', '°')  # Improvement: Clean mangled Unicode
            if not any(key.lower() in header_str for key in ['Cond.', 'pH', 'P Alk', 'P Alkalinity', 'System Type', 'Conductivity', 'uS/cm']):
                logger.debug(f"Table {table_idx} in {filename} does not contain relevant headers, skipping")
                continue

            # Normalize headers
            headers = [re.sub(r'\s+', '_', h.replace('.', '').replace('\n', '').replace('Âµ', 'µ').replace('Â°', '°')).lower().strip() for h in header_row if h.strip()]
            logger.debug(f"Table {table_idx} headers: {headers}")

            for row_idx, row in enumerate(table[1:], 1):
                if not any(row):
                    logger.debug(f"Row {row_idx} in table {table_idx} of {filename} is empty, skipping")
                    continue
                row = [r if r is not None else '' for r in row]
                row += [''] * (len(headers) - len(row))
                d = dict(zip(headers, [r.strip() if r else None for r in row]))
                logger.debug(f"Processing row {row_idx} in table {table_idx}: {d}")

                # Special handling for misplaced values
                glycol = None
                live_atp = None
                for key, val in list(d.items()):
                    if val:
                        val_lower = val.lower()
                        if 'glycol' in val_lower or ('%' in val and 'f' in val_lower):
                            glycol = val
                            if key in [
                                'p_alkalinity_caco3_(mg/l)', 'p_alkalinity',
                                'm_alkalinity_caco3_(mg/l)', 'm_alkalinity',
                                'hardness_caco3_(mg/l)', 'hardness', 'hardnes_s',
                                'ca_(mg/l)', 'ca_caco3_(mg/l)', 'ca', 'calcium',
                                'cond_(µs/cm)', 'cond',
                                'ph',
                                'temp_(°c)', 'temp',
                                'cl_(mg/l)', 'cl', 'chloride',
                                'po4_(mg/l)', 'po4',
                                'no2_(mg/l)', 'no2'
                            ]:
                                d[key] = None  # clear from numeric field
                        if 'atp' in key.lower():
                            if '|' in val:
                                parts = [v.strip() for v in val.split('|')]
                                live_atp = to_float(parts[0]) if parts else None
                            else:
                                live_atp = to_float(val)

                # Determine system name
                system_name = d.get('system_name', '') or d.get('gwt_names', '') or d.get('water_samples', '') or d.get('', '')
                system_type = d.get('system_type', '') or ''
                system_name = (system_type + ' ' + system_name).strip()
                if not system_name:
                    logger.warning(f"No system name extracted for row {row_idx} in table {table_idx} of {filename}, skipping: headers={headers}, row={row}")
                    continue
                logger.debug(f"System name: {system_name}")

                # Extract numerical/text values
                def to_float(val):
                    if val is None or val == '':
                        return None
                    try:
                        cleaned_val = re.sub(r'[^\d.-]', '', val)
                        return float(cleaned_val) if cleaned_val else None
                    except ValueError:
                        return None

                cond = to_float(d.get('cond', None) or d.get('cond_(µs/cm)', None))
                ph = to_float(d.get('ph', None))
                temp = to_float(d.get('temp', None) or d.get('temp_(°c)', None))
                p_alk = to_float(d.get('p_alk', None) or d.get('p_alkalinity_caco3_(mg/l)', None) or d.get('p_alkalinity', None) or d.get('p_alkalinity_caco3_(mg/l)', None))
                m_alk = to_float(d.get('m_alk', None) or d.get('m_alkalinity_caco3_(mg/l)', None) or d.get('m_alkalinity', None) or d.get('m_alkalinity_caco3_(mg/l)', None))
                chloride = to_float(d.get('chloride', None) or d.get('cl_(mg/l)', None) or d.get('cl', None))
                hardness = to_float(d.get('hardness', None) or d.get('hardnes_s', None) or d.get('hardness_caco3_(mg/l)', None))
                calcium = to_float(d.get('calcium', None) or d.get('ca_(mg/l)', None) or d.get('ca_caco3_(mg/l)', None) or d.get('ca', None))
                po4 = to_float(d.get('po4', None) or d.get('po4_(mg/l)', None))
                so3 = to_float(d.get('so2', None) or d.get('so3', None))
                mo = to_float(d.get('mo', None))
                no2 = to_float(d.get('no2', None) or d.get('no2_(mg/l)', None))
                if live_atp is None:
                    live_val = d.get('live_atp', None) or d.get('atp_live_|_dead', None)
                    if live_val:
                        if '|' in live_val:
                            live_atp = to_float(live_val.split('|')[0].strip())
                        else:
                            live_atp = to_float(live_val)
                glycol = glycol or d.get('glycol', None)
                free_chlorine = to_float(d.get('free_chlorine', None) or d.get('free_chlorine_ppm', None))
                total_chlorine = to_float(d.get('total_chlorine', None) or d.get('total_chlorine_ppm', None))
                max_temp = to_float(d.get('max_temp', None))

                # Insert into systems table (Improvement: Check for existing)
                cur.execute("SELECT id FROM systems WHERE report_id = ? AND system_name = ?", (report_id, system_name))
                if cur.fetchone():
                    logger.info(f"Skipping duplicate system {system_name} for report_id {report_id}")
                    continue
                try:
                    cur.execute('''
                    INSERT INTO systems (
                        report_id, system_name, cond, ph, temp, p_alk, m_alk, chloride, hardness, calcium,
                        po4, so3, mo, no2, live_atp, glycol, free_chlorine, total_chlorine, max_temp, comment
                    ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    ''', (
                        report_id, system_name, cond, ph, temp, p_alk, m_alk, chloride, hardness, calcium,
                        po4, so3, mo, no2, live_atp, glycol, free_chlorine, total_chlorine, max_temp, comment
                    ))
                    conn.commit()
                    logger.info(f"Inserted system for report_id {report_id}: {system_name}")
                except Exception as e:
                    logger.error(f"Error inserting system {system_name} for {filename}: {e}")
                    continue

    # Close the connection
    try:
        conn.close()
        logger.info("Database connection closed")
    except Exception as e:
        logger.error(f"Error closing database connection: {e}")

    logger.info("Data extraction and insertion completed")

if __name__ == "__main__":
    # For testing the function directly
    script_dir = os.path.dirname(os.path.abspath(__file__))
    reports_path = os.path.join(script_dir, 'reports')
    db_path = os.path.join(script_dir, 'reports.db')  # Place reports.db in data_processing
    process_reports(reports_path, db_path)